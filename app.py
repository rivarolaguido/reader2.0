import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from urllib.parse import urlparse

# ─── Page Config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="URL Text Transcriber",
    page_icon="📄",
    layout="centered"
)

# ─── Custom CSS ─────────────────────────────────────────────────────────────
st.markdown("""
    <style>
        /* Hide Streamlit default UI elements */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}

        /* General */
        body {
            background-color: #f4f6f9;
            font-family: 'Segoe UI', sans-serif;
        }
        .block-container {
            max-width: 700px;
            padding-top: 60px;
        }

        /* Title */
        .app-title {
            text-align: center;
            font-size: 2.2rem;
            font-weight: 700;
            color: #1a73e8;
            margin-bottom: 5px;
        }
        .app-subtitle {
            text-align: center;
            font-size: 1rem;
            color: #666;
            margin-bottom: 35px;
        }

        /* Input */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea {
            border-radius: 10px;
            border: 1.5px solid #d0d7de;
            padding: 12px;
            font-size: 15px;
            background-color: #ffffff;
        }
        .stTextInput > div > div > input:focus,
        .stTextArea > div > div > textarea:focus {
            border-color: #1a73e8;
            box-shadow: 0 0 0 3px rgba(26,115,232,0.15);
        }

        /* Radio */
        .stRadio > div {
            justify-content: center;
            gap: 20px;
        }

        /* Button */
        .stButton > button {
            background-color: #1a73e8;
            color: white;
            font-size: 16px;
            font-weight: 600;
            padding: 12px 0;
            border-radius: 10px;
            border: none;
            width: 100%;
            transition: background-color 0.2s ease;
        }
        .stButton > button:hover {
            background-color: #1558b0;
        }

        /* Download Button */
        .stDownloadButton > button {
            background-color: #34a853;
            color: white;
            font-size: 15px;
            font-weight: 600;
            padding: 12px 0;
            border-radius: 10px;
            border: none;
            width: 100%;
            transition: background-color 0.2s ease;
        }
        .stDownloadButton > button:hover {
            background-color: #2d8f47;
        }

        /* Alert boxes */
        .stAlert {
            border-radius: 10px;
        }

        /* Divider */
        hr {
            border: none;
            border-top: 1px solid #e0e0e0;
            margin: 25px 0;
        }

        /* Footer note */
        .footer-note {
            text-align: center;
            font-size: 0.8rem;
            color: #aaa;
            margin-top: 30px;
        }
    </style>
""", unsafe_allow_html=True)


# ─── Helper: Validate URL ────────────────────────────────────────────────────
def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme in ("http", "https"), result.netloc])
    except:
        return False


# ─── Helper: Scrape Text Only (No UI Elements) ──────────────────────────────
def scrape_text(url):
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/114.0.0.0 Safari/537.36"
        )
    }

    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "lxml")

    # ── Remove ALL non-text elements ──
    remove_tags = [
        "script", "style", "img", "nav", "footer", "header",
        "aside", "iframe", "noscript", "svg", "figure", "figcaption",
        "form", "input", "button", "select", "textarea", "label",
        "menu", "menuitem", "dialog", "canvas", "video", "audio",
        "map", "area", "object", "embed", "param", "source", "track",
        "picture", "link", "meta", "head"
    ]
    for tag in soup(remove_tags):
        tag.decompose()

    # ── Remove elements by common UI class/id keywords ──
    ui_keywords = [
        "nav", "navbar", "navigation", "menu", "sidebar", "side-bar",
        "header", "footer", "banner", "cookie", "popup", "modal",
        "advertisement", "ad", "ads", "promo", "social", "share",
        "comment", "comments", "related", "recommended", "subscribe",
        "newsletter", "breadcrumb", "pagination", "widget", "toolbar",
        "tooltip", "dropdown", "overlay", "notification", "alert-bar"
    ]
    for element in soup.find_all(True):
        el_id = element.get("id", "").lower()
        el_class = " ".join(element.get("class", [])).lower()
        if any(kw in el_id or kw in el_class for kw in ui_keywords):
            element.decompose()

    # ── Get page title ──
    title_tag = soup.find("title")
    page_title = title_tag.get_text(strip=True) if title_tag else "Untitled Page"

    # ── Extract only meaningful text content ──
    content = []
    seen_texts = set()  # Avoid duplicate lines

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote"]):
        text = element.get_text(separator=" ", strip=True)
        text = re.sub(r'\s+', ' ', text).strip()

        # Skip if: empty, too short, already seen, or looks like UI text
        if (
            not text
            or len(text) < 15
            or text in seen_texts
            or re.match(r'^[\W\d]+$', text)  # Only symbols/numbers
        ):
            continue

        seen_texts.add(text)
        content.append({"tag": element.name, "text": text})

    return page_title, content


# ─── Helper: Build Word Document ─────────────────────────────────────────────
def build_docx(page_title, content, url):
    doc = Document()

    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(80)
        section.right_margin = Pt(80)

    # ── Document Title ──
    title_para = doc.add_heading(page_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    # ── Source URL ──
    source_para = doc.add_paragraph()
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_run = source_para.add_run(f"Source: {url}")
    source_run.font.size = Pt(9)
    source_run.font.italic = True
    source_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()  # Spacer

    # ── Heading level map ──
    heading_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

    # ── Write Content ──
    for item in content:
        tag = item["tag"]
        text = item["text"]

        if tag in heading_map:
            doc.add_heading(text, level=heading_map[tag])

        elif tag == "p":
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(11)

        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")

        elif tag == "blockquote":
            para = doc.add_paragraph(text, style="Intense Quote")

    # ── Footer ──
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Transcribed by URL Transcriber  |  {url}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_para.runs:
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save to buffer ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Helper: Generate filename ───────────────────────────────────────────────
def url_to_filename(url):
    parsed = urlparse(url)
    name = parsed.path.strip("/").replace("/", "_") or parsed.netloc
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    return f"{name[:60]}.docx" if name else "transcription.docx"


# ════════════════════════════════════════════════════════════════════════════
#  MAIN APP UI
# ════════════════════════════════════════════════════════════════════════════

# ── Header ──
st.markdown('<div class="app-title">📄 URL Text Transcriber</div>', unsafe_allow_html=True)
st.markdown('<div class="app-subtitle">Paste a URL and download the text as a Word document — clean, no images.</div>', unsafe_allow_html=True)

st.markdown("<hr>", unsafe_allow_html=True)

# ── Mode Toggle ──
mode = st.radio("", ["Single URL", "Multiple URLs"], horizontal=True, label_visibility="collapsed")

urls = []

if mode == "Single URL":
    url_input = st.text_input("", placeholder="🔗  https://example.com/article", label_visibility="collapsed")
    if url_input:
        urls = [url_input.strip()]
else:
    url_text = st.text_area(
        "",
        placeholder="🔗  https://example.com/article-1\n🔗  https://example.com/article-2\n🔗  https://example.com/article-3",
        height=160,
        label_visibility="collapsed"
    )
    if url_text:
        urls = [u.strip() for u in url_text.strip().splitlines() if u.strip()]

st.markdown("<br>", unsafe_allow_html=True)

# ── Transcribe Button ──
if st.button("✨ Transcribe & Download"):

    if not urls:
        st.warning("⚠️ Please enter at least one URL.")

    else:
        invalid = [u for u in urls if not is_valid_url(u)]
        if invalid:
            st.error(f"❌ Invalid URL(s): {', '.join(invalid)}")

        # ── Single URL ──
        elif len(urls) == 1:
            url = urls[0]
            with st.spinner("Extracting text content..."):
                try:
                    page_title, content = scrape_text(url)

                    if not content:
                        st.error("❌ No readable text found on that page.")
                    else:
                        docx_buffer = build_docx(page_title, content, url)
                        filename = url_to_filename(url)

                        st.success(f"✅ **{len(content)} text blocks** extracted from: _{page_title}_")

                        st.download_button(
                            label="📥 Download Word Document",
                            data=docx_buffer,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                except requests.exceptions.ConnectionError:
                    st.error("❌ Could not connect. Check the URL or your internet connection.")
                except requests.exceptions.Timeout:
                    st.error("❌ Request timed out. The site may be too slow.")
                except requests.exceptions.HTTPError as e:
                    st.error(f"❌ HTTP Error: {e}")
                except Exception as e:
                    st.error(f"❌ Unexpected error: {e}")

        # ── Multiple URLs ──
        else:
            combined_doc = Document()
            combined_doc.add_heading("Combined URL Transcription", level=0)
            combined_doc.add_paragraph()

            success_count = 0
            progress = st.progress(0)
            status = st.empty()

            heading_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

            for i, url in enumerate(urls):
                status.text(f"Processing {i + 1}/{len(urls)}: {url}")
                try:
                    page_title, content = scrape_text(url)
                    if content:
                        combined_doc.add_heading(page_title, level=1)
                        src = combined_doc.add_paragraph(f"Source: {url}")
                        src.runs[0].font.italic = True
                        src.runs[0].font.size = Pt(9)
                        combined_doc.add_paragraph()

                        for item in content:
                            tag = item["tag"]
                            text = item["text"]
                            if tag in heading_map:
                                combined_doc.add_heading(text, level=heading_map[tag])
                            elif tag == "p":
                                combined_doc.add_paragraph(text)
                            elif tag == "li":
                                combined_doc.add_paragraph(text, style="List Bullet")
                            elif tag == "blockquote":
                                combined_doc.add_paragraph(text, style="Intense Quote")

                        combined_doc.add_page_break()
                        success_count += 1

                except Exception as e:
                    combined_doc.add_paragraph(f"⚠️ Failed to scrape {url}: {e}")

                progress.progress((i + 1) / len(urls))

            status.empty()
            progress.empty()

            buffer = io.BytesIO()
            combined_doc.save(buffer)
            buffer.seek(0)

            st.success(f"✅ Successfully transcribed **{success_count}/{len(urls)}** URLs.")
            st.download_button(
                label="📥 Download Combined Word Document",
                data=buffer,
                file_name="combined_transcription.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown('<div class="footer-note">⚠️ Some websites may block automated requests. Results may vary.</div>', unsafe_allow_html=True)
